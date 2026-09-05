#!/usr/bin/env python3
"""Build the Scientific Data v7 working manuscript from its Markdown source.

This script writes only to the new v7 manuscript directory.
"""

from __future__ import annotations

import re
import hashlib
import json
import platform
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "OvCAN_Scientific_Data_draft_v7.md"
OUTPUT = ROOT / "OvCAN_Scientific_Data_draft_v7.docx"

BLUE = "000000"
DARK_BLUE = "000000"
TEXT = "000000"
MUTED = "666666"
PALE_BLUE = "F4F6F9"
PLACEHOLDER_FILL = "FFF2CC"
PLACEHOLDER_TEXT = "9C5700"
GRID = "D9D9D9"
WHITE = "FFFFFF"

INLINE_PATTERN = re.compile(
    r"(\[[A-Z][^\]]+\]|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\^[^^]+\^|https?://\S+)"
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_layout(table, widths: list[int], description: str) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), GRID)
        borders.append(border)
    tbl_pr.append(borders)

    tbl_desc = OxmlElement("w:tblDescription")
    tbl_desc.set(qn("w:val"), description)
    tbl_pr.append(tbl_desc)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_flags(paragraph, *, keep_next=False, keep_lines=False, widow=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for tag, enabled in (
        ("keepNext", keep_next),
        ("keepLines", keep_lines),
        ("widowControl", widow),
    ):
        node = p_pr.find(qn(f"w:{tag}"))
        if enabled and node is None:
            node = OxmlElement(f"w:{tag}")
            p_pr.append(node)
        elif enabled and node is not None:
            node.set(qn("w:val"), "1")


def set_run_font(run, name="Arial", size=11, color=TEXT) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_inline(paragraph, text: str, *, size=11, color=TEXT) -> None:
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, color=color)

        token = match.group(0)
        run_text = token
        italic = False
        bold = False
        superscript = False
        code = False
        placeholder = False
        hyperlink = False

        if token.startswith("["):
            placeholder = True
        elif token.startswith("**"):
            run_text = token[2:-2]
            bold = True
        elif token.startswith("*"):
            run_text = token[1:-1]
            italic = True
        elif token.startswith("`"):
            run_text = token[1:-1]
            code = True
        elif token.startswith("^"):
            run_text = token[1:-1]
            superscript = True
        elif token.startswith("http"):
            hyperlink = True

        run = paragraph.add_run(run_text)
        set_run_font(
            run,
            name="Arial" if code else "Arial",
            size=9 if (code or superscript) else size,
            color=PLACEHOLDER_TEXT if placeholder else (BLUE if hyperlink else color),
        )
        run.italic = italic
        run.bold = bold
        run.font.superscript = superscript
        run.font.underline = hyperlink
        if placeholder:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PLACEHOLDER_FILL)
            run._element.get_or_add_rPr().append(shd)
        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    # Remove decorative paragraph rules inherited from the runtime template.
    for border in styles.element.xpath(".//w:pBdr"):
        border.getparent().remove(border)

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True

    for style_name, size, color, before, after in (
        ("Heading 1", 14, BLUE, 14, 6),
        ("Heading 2", 12, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(10)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    subtitle.paragraph_format.space_after = Pt(10)

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(10)
    caption.font.bold = True
    caption.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption.paragraph_format.space_before = Pt(10)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    list_number = styles["List Number"]
    list_number.font.name = "Arial"
    list_number.font.size = Pt(9.5)
    list_number.font.color.rgb = RGBColor.from_string(TEXT)
    list_number._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    list_number.paragraph_format.left_indent = Inches(0.30)
    list_number.paragraph_format.first_line_indent = Inches(-0.30)
    list_number.paragraph_format.line_spacing = 1.08
    list_number.paragraph_format.space_after = Pt(4)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("")
    set_run_font(run, size=8, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    set_run_font(run, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_table(doc: Document, rows: list[list[str]], table_number: int) -> None:
    widths = (
        [2390, 970, 1040, 650, 1040, 1120, 1090, 1060]
        if table_number == 1
        else [1450, 3600, 1750, 2560]
    )
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    description = (
        "Numbers of ovarian cancer models and patients represented by each assay and histotype."
        if table_number == 1
        else "Principal deposited data records, their contents, units and intended repositories."
    )
    set_table_layout(table, widths, description)
    set_repeat_table_header(table.rows[0])

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            if row_idx == 0:
                set_cell_shading(cell, "E8ECF0")
            elif row_idx % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if col_idx == 0 or table_number == 2
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            if table_number == 1:
                set_cell_margins(cell, start=70, end=70)
            add_inline(p, value, size=9, color=TEXT)
            for run in p.runs:
                if row_idx == 0:
                    run.bold = True
            set_paragraph_flags(p, keep_lines=True)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


def parse_markdown(doc: Document, source: str) -> None:
    lines = source.splitlines()
    idx = 0
    current_major = ""
    table_number = 0

    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line:
            idx += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, line[2:].strip(), size=20, color=DARK_BLUE)
            set_paragraph_flags(p, keep_next=True, keep_lines=True)
            idx += 1
            continue

        if line.startswith("## "):
            current_major = line[3:].strip()
            if current_major in {"Tables", "Figure legends"}:
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, current_major, size=14, color=BLUE)
            set_paragraph_flags(p, keep_next=True, keep_lines=True)
            idx += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            style = "Caption" if current_major in {"Tables", "Figure legends"} else "Heading 2"
            p = doc.add_paragraph(style=style)
            add_inline(p, text, size=10 if style == "Caption" else 12, color=DARK_BLUE if style == "Caption" else BLUE)
            set_paragraph_flags(p, keep_next=True, keep_lines=True)
            if current_major == "Tables":
                match = re.match(r"Table\s+(\d+)", text)
                table_number = int(match.group(1)) if match else table_number + 1
            idx += 1
            continue

        if line.startswith("|"):
            raw_rows = []
            while idx < len(lines) and lines[idx].lstrip().startswith("|"):
                raw_rows.append(lines[idx].strip())
                idx += 1
            rows = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in raw_rows
                if not re.fullmatch(r"\|?[\s:|-]+\|?", row)
            ]
            add_table(doc, rows, table_number)
            continue

        if current_major == "References" and re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, text, size=9.5)
            set_paragraph_flags(p, keep_lines=True)
            idx += 1
            continue

        if line.startswith("*Working draft"):
            p = doc.add_paragraph(style="Subtitle")
            add_inline(p, line[1:-1], size=10, color=MUTED)
            idx += 1
            continue

        p = doc.add_paragraph(style="Normal")
        if current_major == "Figure legends":
            p.paragraph_format.line_spacing = 1.1
            add_inline(p, line, size=10)
            set_paragraph_flags(p, keep_lines=True, widow=True)
        else:
            add_inline(p, line)
            set_paragraph_flags(p, keep_lines=line.startswith("["), widow=True)
        idx += 1


def add_document_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.title = SOURCE.read_text(encoding="utf-8").splitlines()[0][2:]
    props.subject = "Scientific Data working manuscript"
    props.author = "OvCAN study team"
    props.comments = "Version 7; explicit TODO markers link to author_confirmation.md."


def add_figures(doc: Document) -> None:
    figures = ROOT.parent / "figures"
    for number in range(1, 5):
        matches = sorted(figures.glob(f"Figure{number}*.png")) + sorted(figures.glob(f"figure{number}*.png")) + sorted(figures.glob(f"fig{number}*.png"))
        if len(matches) != 1:
            raise ValueError(f"Expected one PNG for figure {number}; found {matches}")
        doc.add_page_break()
        heading = doc.add_paragraph(style="Heading 1")
        heading.add_run(f"Figure {number}")
        heading.paragraph_format.space_before = Pt(0)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        shape = run.add_picture(str(matches[0]), width=Inches(6.5))
        shape._inline.docPr.set("descr", f"Figure {number}. See the Figure legends section for a panel-by-panel description.")


def main() -> None:
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_document_metadata(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    add_figures(doc)
    for border in doc.element.xpath(".//w:pBdr"):
        border.getparent().remove(border)
    doc.save(OUTPUT)
    inputs = [SOURCE, Path(__file__).resolve(), *[ROOT.parent / "figures" / f"fig{i}.png" for i in range(1, 5)]]
    manifest = {
        "python": platform.python_version(),
        "inputs": [{"path": str(p.relative_to(ROOT.parent.parent.parent)), "sha256": hashlib.sha256(p.read_bytes()).hexdigest()} for p in inputs],
        "output": {"path": OUTPUT.name, "sha256": hashlib.sha256(OUTPUT.read_bytes()).hexdigest()},
    }
    (ROOT / "build_manifest.json").write_text(json.dumps(manifest, indent=2) + "\n")
    print(OUTPUT)


if __name__ == "__main__":
    main()
